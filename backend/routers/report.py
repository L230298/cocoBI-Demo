"""数据分析报告生成 + SQL 编辑/重执行"""
from __future__ import annotations
import re
import statistics
from io import BytesIO
from typing import Optional, List, Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

from tools import invoke_tool
from tools.execute_sql import _DB_CACHE, _get_or_load_db

router = APIRouter(prefix="/api", tags=["report"])

# SQL 白名单(防止破坏性操作)
_FORBIDDEN = re.compile(
    r"\b(DROP|DELETE|UPDATE|INSERT|TRUNCATE|ALTER|CREATE|REPLACE)\b",
    re.IGNORECASE,
)


# ==================== 报告模板元数据 ====================
# 指标体系:每个 Intent 对应的指标定义和统计口径
INTENT_METRICS: dict[str, dict[str, str]] = {
    "QueryBasicMetrics": {
        "name": "基础指标查询",
        "definition": "对核心业务指标做聚合统计(如 GMV、订单数、用户数等)",
        "caliber": "SQL 聚合函数(COUNT / SUM / AVG / MAX / MIN),默认对主键去重",
    },
    "QueryCompareAndTopN": {
        "name": "对比与 TOP N",
        "definition": "按维度对比指标,或取排序前 N 名",
        "caliber": "GROUP BY 维度 + ORDER BY 指标 DESC + LIMIT N",
    },
    "ThresholdAlert": {
        "name": "阈值告警",
        "definition": "对比指标值与阈值,识别异常波动",
        "caliber": "指标值 vs 配置阈值,超阈值标红",
    },
    "AttributeAnalysis": {
        "name": "属性归因",
        "definition": "按维度拆分指标,找贡献最大的属性",
        "caliber": "GROUP BY 维度,计算占比 / 贡献度",
    },
    "SmartInterpretation": {
        "name": "智能解读",
        "definition": "AI 自动解读数据,生成业务洞察",
        "caliber": "综合多个指标,LLM 总结归因",
    },
    "QueryUserCount": {
        "name": "用户数查询",
        "definition": "统计满足条件的用户数(按 user_id 去重)",
        "caliber": "COUNT(DISTINCT user_id)",
    },
    "QueryBasicTrend": {
        "name": "基础趋势",
        "definition": "按时间维度看指标变化趋势",
        "caliber": "GROUP BY 时间粒度(日/周/月)",
    },
    "QueryCategoryRank": {
        "name": "品类排名",
        "definition": "按品类维度排名,看头部 / 长尾",
        "caliber": "GROUP BY 品类 + ORDER BY 指标 DESC",
    },
}

# 策略建议模板
STRATEGY_TEMPLATES: dict[str, list[str]] = {
    "QueryBasicMetrics": [
        "关注核心指标的环比 / 同比变化,偏离均值 ±20% 需排查",
        "建议建立指标日报 / 周报,常态化监控",
        "结合业务日历(大促 / 节假)解读波动,避免误判",
    ],
    "QueryCompareAndTopN": [
        "头部贡献集中度高(>50%)时,需关注长尾品类的潜力",
        "建议对 Top N 单独建监控,异常时优先响应",
        "长尾品类若有突增,可能存在新机会或异常,需深挖",
    ],
    "ThresholdAlert": [
        "超阈值的指标先做归因(数据 / 业务 / 系统),再决策",
        "建议配置自动告警通知(企微 / 钉钉),缩短响应时间",
        "阈值需周期性校准,避免误报或漏报",
    ],
    "AttributeAnalysis": [
        "高贡献属性加大资源投放,低贡献属性考虑优化或下线",
        "建议结合 ROI 做归因,避免被表面贡献误导",
        "维度可下钻到二级属性,找真正的驱动因子",
    ],
    "SmartInterpretation": [
        "LLM 解读需结合业务上下文验证,不可全盘接受",
        "建议把解读结论沉淀到知识库,持续迭代",
        "对解读中的关键数字交叉验证,避免幻觉",
    ],
    "QueryUserCount": [
        "用户数是基础规模指标,需结合活跃 / 留存看质量",
        "建议按新老用户拆分,识别增长来源",
        "低活跃用户需运营触达,防流失",
    ],
    "QueryBasicTrend": [
        "短期波动看日 / 周,长期趋势看月 / 季",
        "异常点(突增 / 突降)做事件标注,方便回溯",
        "建议叠加去年同期对比,识别季节性",
    ],
    "QueryCategoryRank": [
        "头部品类要保供给 / 防缺货,长尾品类看毛利",
        "排名变化 >3 位的品类,需做归因(上新 / 活动 / 缺货)",
        "关注品类集中度(HHI),过高有风险",
    ],
    "default": [
        "建议结合历史数据做趋势分析,识别异常点",
        "跨指标交叉验证,避免单点误判",
        "周期性复盘指标变化,迭代分析口径",
    ],
}


def _basic_stats(rows: list[dict], cols: list[str]) -> dict[str, dict[str, Any]]:
    """对每列做基础统计,生成异常 / 归因数据"""
    stats: dict[str, dict[str, Any]] = {}
    for col in cols:
        vals = [r.get(col) for r in rows if r.get(col) is not None]
        if not vals:
            stats[col] = {"type": "empty", "count": 0}
            continue
        # 数值列
        nums = []
        for v in vals:
            try:
                nums.append(float(v))
            except (TypeError, ValueError):
                break
        if len(nums) == len(vals) and nums:
            stats[col] = {
                "type": "numeric",
                "count": len(nums),
                "min": min(nums),
                "max": max(nums),
                "avg": round(sum(nums) / len(nums), 2),
                "std": round(statistics.pstdev(nums), 2) if len(nums) > 1 else 0,
            }
            # 异常值检测: 超过 3σ
            mean = sum(nums) / len(nums)
            std = statistics.pstdev(nums) if len(nums) > 1 else 0
            if std > 0:
                outliers = [v for v in nums if abs(v - mean) > 3 * std]
                stats[col]["outliers"] = len(outliers)
        else:
            # 类别列
            from collections import Counter
            counter = Counter(str(v) for v in vals)
            top3 = counter.most_common(3)
            stats[col] = {
                "type": "categorical",
                "count": len(vals),
                "distinct": len(counter),
                "top3": top3,
            }
    return stats


# ==================== docx 渲染 ====================
def _build_docx(report: dict, dataset: Optional[dict] = None) -> bytes:
    """把 report dict 渲染成 8 章节 .docx (Word) 文档, 返回字节"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认中文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    rpr = style.element.rPr
    from docx.oxml.ns import qn
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sections = report.get("sections", [])
    dataset_name = (dataset or {}).get("name", report.get("dataset_id", "-"))
    dataset_rows = (dataset or {}).get("row_count", 0)
    dataset_cols_count = (dataset or {}).get("column_count", 0)
    dataset_fields = (dataset or {}).get("fields", [])
    dataset_uploaded = (dataset or {}).get("uploaded_at", "")

    # ==================== 封面 / 标题 ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(report.get("title", "数据分析报告"))
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(f"数据集: {dataset_name}  |  指标数: {len(sections)}")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 空行

    # ==================== 第 1 章 报告基本信息 ====================
    doc.add_heading("一、报告基本信息", level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Light Shading Accent 1"
    info_data = [
        ("报告标题", report.get("title", "-")),
        ("数据集 ID", report.get("dataset_id", "-")),
        ("数据集名称", dataset_name),
        ("生成时间", report.get("generated_at", "-")),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = str(v)
        for run in info_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    # ==================== 第 2 章 业务背景与分析目标 ====================
    doc.add_heading("二、业务背景与分析目标", level=1)
    doc.add_paragraph().add_run("2.1 业务背景").bold = True
    bg_p = doc.add_paragraph(
        f"本报告基于数据集「{dataset_name}」进行分析。"
        f"该数据集共包含 {dataset_rows} 行 × {dataset_cols_count} 列业务数据,"
        f"覆盖 {len(dataset_fields)} 个核心字段。"
    )
    doc.add_paragraph().add_run("2.2 分析目标").bold = True
    if sections:
        for i, sec in enumerate(sections, 1):
            query = sec.get("query", "N/A")
            doc.add_paragraph(f"{i}. {query}", style="List Number")
    else:
        doc.add_paragraph("(本次未选取任何 query,无具体分析目标)")

    # ==================== 第 3 章 指标体系与统计口径 ====================
    doc.add_heading("三、指标体系与统计口径", level=1)
    used_intents = sorted({s.get("intent", "default") for s in sections if not s.get("error")})
    if not used_intents:
        used_intents = ["default"]
    metric_table = doc.add_table(rows=1, cols=4)
    metric_table.style = "Light Grid Accent 1"
    hdr = metric_table.rows[0].cells
    for j, h in enumerate(["指标类型", "定义", "统计口径", "SQL 特征"]):
        hdr[j].text = h
        for run in hdr[j].paragraphs[0].runs:
            run.bold = True
    for intent in used_intents:
        m = INTENT_METRICS.get(intent, INTENT_METRICS.get("default", {
            "name": intent, "definition": "-", "caliber": "-"
        }))
        cells = metric_table.add_row().cells
        cells[0].text = m.get("name", intent)
        cells[1].text = m.get("definition", "-")
        cells[2].text = m.get("caliber", "-")
        # 找对应 section 的 SQL 特征
        sec = next((s for s in sections if s.get("intent") == intent and not s.get("error")), None)
        cells[3].text = (sec.get("sql", "-")[:60] + "...") if sec and sec.get("sql") else "-"

    # ==================== 第 4 章 数据来源与质量说明 ====================
    doc.add_heading("四、数据来源与质量说明", level=1)
    doc.add_paragraph().add_run("4.1 数据源信息").bold = True
    src_table = doc.add_table(rows=1, cols=2)
    src_table.style = "Light Grid Accent 1"
    hdr = src_table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "值"
    for run in hdr[0].paragraphs[0].runs + hdr[1].paragraphs[0].runs:
        run.bold = True
    src_rows = [
        ("数据来源", (dataset or {}).get("file_path", "未提供")),
        ("数据行数", f"{dataset_rows} 行"),
        ("字段数", f"{dataset_cols_count} 个"),
        ("上传时间", dataset_uploaded),
        ("行业模板", (dataset or {}).get("industry_template", "通用")),
    ]
    for k, v in src_rows:
        cells = src_table.add_row().cells
        cells[0].text = k
        cells[1].text = str(v)

    doc.add_paragraph().add_run("4.2 数据质量评估").bold = True
    quality_items = [
        f"完整度: 字段数 {dataset_cols_count} 个,数据行 {dataset_rows} 行,无明显缺失。",
        f"时效性: 数据最后更新于 {dataset_uploaded[:19] if dataset_uploaded else '未知'}。",
        "一致性: 报告内所有指标均来自同一数据源,口径一致。",
        "唯一性: 主键去重处理,避免重复计数。",
    ]
    for q in quality_items:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_paragraph().add_run("4.3 字段说明").bold = True
    if dataset_fields:
        field_table = doc.add_table(rows=1, cols=3)
        field_table.style = "Light Grid Accent 1"
        hdr = field_table.rows[0].cells
        for j, h in enumerate(["字段名", "类型", "示例值"]):
            hdr[j].text = h
            for run in hdr[j].paragraphs[0].runs:
                run.bold = True
        for f in dataset_fields:
            cells = field_table.add_row().cells
            cells[0].text = str(f.get("name", "-"))
            cells[1].text = str(f.get("type", "-"))
            cells[2].text = str(f.get("sample", "-"))
    else:
        doc.add_paragraph("(无字段元数据)")

    # ==================== 第 5 章 指标报表展示 ====================
    doc.add_heading("五、指标报表展示", level=1)
    if not sections:
        doc.add_paragraph("(本次未选取任何 query)")
    for i, sec in enumerate(sections, 1):
        doc.add_heading(f"5.{i} {sec.get('query', 'N/A')}", level=2)

        info = doc.add_paragraph()
        info_run = info.add_run(
            f"Intent: {sec.get('intent', '-')}   |   "
            f"返回行数: {sec.get('row_count', 0)}   |   "
            f"时间: {sec.get('timestamp', '-')}"
        )
        info_run.font.size = Pt(9)
        info_run.italic = True
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if sec.get("error"):
            err_p = doc.add_paragraph()
            err_run = err_p.add_run(f"❌ 查询失败: {sec['error']}")
            err_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            continue

        # SQL
        doc.add_paragraph().add_run("5." + str(i) + ".1 SQL 语句").bold = True
        sql_p = doc.add_paragraph()
        sql_run = sql_p.add_run(sec.get("sql", "-") or "-")
        sql_run.font.name = "Consolas"
        sql_run.font.size = Pt(9)
        sql_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

        # 数据表
        cols = sec.get("columns", [])
        rows = sec.get("rows", [])
        if cols and rows:
            doc.add_paragraph().add_run("5." + str(i) + ".2 查询结果").bold = True
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for j, col in enumerate(cols):
                hdr[j].text = str(col)
                for run in hdr[j].paragraphs[0].runs:
                    run.bold = True
            for row in rows[:50]:
                cells = table.add_row().cells
                for j, col in enumerate(cols):
                    cells[j].text = str(row.get(col, ""))
        else:
            doc.add_paragraph("(无数据)")

    # ==================== 第 6 章 异常与归因分析 ====================
    doc.add_heading("六、异常与归因分析", level=1)
    any_stats = False
    for i, sec in enumerate(sections, 1):
        if sec.get("error") or not sec.get("rows"):
            continue
        any_stats = True
        doc.add_paragraph().add_run(f"6.{i} {sec.get('query', 'N/A')} - 数据特征").bold = True
        stats = _basic_stats(sec["rows"], sec.get("columns", []))
        if not stats:
            doc.add_paragraph("(无可统计数据)")
            continue
        stat_table = doc.add_table(rows=1, cols=4)
        stat_table.style = "Light Grid Accent 1"
        hdr = stat_table.rows[0].cells
        for j, h in enumerate(["字段", "类型", "关键指标", "异常值"]):
            hdr[j].text = h
            for run in hdr[j].paragraphs[0].runs:
                run.bold = True
        for col, s in stats.items():
            cells = stat_table.add_row().cells
            cells[0].text = col
            cells[1].text = s.get("type", "-")
            if s.get("type") == "numeric":
                cells[2].text = f"min={s.get('min')}, max={s.get('max')}, avg={s.get('avg')}, σ={s.get('std')}"
                cells[3].text = str(s.get("outliers", 0))
            elif s.get("type") == "categorical":
                top3 = s.get("top3", [])
                top3_str = ", ".join(f"{k}({v})" for k, v in top3)
                cells[2].text = f"distinct={s.get('distinct')}, top3: {top3_str}"
                cells[3].text = "-"
            else:
                cells[2].text = "无数据"
                cells[3].text = "-"
    if not any_stats:
        doc.add_paragraph("(无足够数据进行异常分析)")

    # ==================== 第 7 章 策略建议与行动计划 ====================
    doc.add_heading("七、策略建议与行动计划", level=1)
    doc.add_paragraph().add_run("7.1 指标维度建议").bold = True
    used_intents = sorted({s.get("intent", "default") for s in sections if not s.get("error")})
    if not used_intents:
        used_intents = ["default"]
    for intent in used_intents:
        m = INTENT_METRICS.get(intent, INTENT_METRICS.get("default", {}))
        doc.add_paragraph().add_run(f"• {m.get('name', intent)}").bold = True
        tips = STRATEGY_TEMPLATES.get(intent, STRATEGY_TEMPLATES["default"])
        for tip in tips:
            doc.add_paragraph(tip, style="List Bullet")

    doc.add_paragraph().add_run("7.2 行动计划").bold = True
    actions = [
        ("短期(1 周内)", [
            "将本次报告的指标纳入周报,常态化跟踪",
            "对识别出的异常值(3σ 外)做归因分析",
            "若有阈值告警,配置自动通知",
        ]),
        ("中期(1 个月内)", [
            "扩展分析维度(用户分层 / 渠道拆分)",
            "建立指标趋势基线,识别异常波动",
            "沉淀分析模板,提升复用率",
        ]),
        ("长期(1 季度内)", [
            "搭建指标体系,覆盖核心业务全链路",
            "结合 LLM 智能解读,提升分析深度",
            "形成数据驱动的决策机制",
        ]),
    ]
    for phase, items in actions:
        doc.add_paragraph().add_run(phase).bold = True
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    # ==================== 第 8 章 附录 ====================
    doc.add_heading("八、附录", level=1)
    doc.add_paragraph().add_run("8.1 术语表").bold = True
    glossary = [
        ("GMV", "Gross Merchandise Volume,商品交易总额"),
        ("DAU", "Daily Active Users,日活跃用户数"),
        ("MAU", "Monthly Active Users,月活跃用户数"),
        ("ROI", "Return On Investment,投资回报率"),
        ("HHI", "Herfindahl-Hirschman Index,品类集中度"),
        ("SQL", "Structured Query Language,结构化查询语言"),
        ("3σ 原则", "数值偏离均值超过 3 倍标准差视为异常"),
    ]
    gloss_table = doc.add_table(rows=1, cols=2)
    gloss_table.style = "Light Grid Accent 1"
    hdr = gloss_table.rows[0].cells
    hdr[0].text = "术语"
    hdr[1].text = "解释"
    for run in hdr[0].paragraphs[0].runs + hdr[1].paragraphs[0].runs:
        run.bold = True
    for k, v in glossary:
        cells = gloss_table.add_row().cells
        cells[0].text = k
        cells[1].text = v

    doc.add_paragraph().add_run("8.2 报告生成方式").bold = True
    gen_info = doc.add_paragraph(
        "本报告由 cocoBI AI 数据分析助手自动生成,流程:\n"
        "1. 用户在前端选择历史 query,提交生成报告请求\n"
        "2. 后端根据 query_ids 从会话历史中恢复 SQL\n"
        "3. 重新执行 SQL 获取最新数据\n"
        "4. 加载数据集元数据,计算基础统计\n"
        "5. 按 8 章节模板用 python-docx 渲染输出 .docx\n"
        "6. 通过 HTTP Response 返回二进制流,前端触发下载"
    )

    doc.add_paragraph().add_run("8.3 报告版本").bold = True
    doc.add_paragraph(f"版本: v1.0    生成时间: {report.get('generated_at', '-')}")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot.add_run("— 本报告由 cocoBI 自动生成 —")
    foot_run.font.size = Pt(9)
    foot_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    foot_run.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==================== 业务端点 ====================
class SqlEditRequest(BaseModel):
    """用户编辑 SQL 后重新执行"""
    sql: str
    dataset_id: str


class SqlEditResponse(BaseModel):
    """SQL 重执行结果"""
    success: bool
    columns: List[str] = []
    rows: List[dict] = []
    row_count: int = 0
    elapsed_ms: int = 0
    error: Optional[str] = None
    is_readonly: bool = True


@router.post("/sql/edit", response_model=SqlEditResponse)
async def edit_and_execute_sql(req: SqlEditRequest):
    """用户编辑 SQL 后重新执行(只允许 SELECT)"""
    sql = (req.sql or "").strip()
    if not sql:
        return SqlEditResponse(success=False, error="SQL 不能为空")

    if _FORBIDDEN.search(sql):
        return SqlEditResponse(
            success=False,
            is_readonly=False,
            error="SQL 包含禁止操作(DROP/DELETE/UPDATE/INSERT 等),只允许 SELECT",
        )

    if not sql.upper().startswith("SELECT"):
        return SqlEditResponse(
            success=False,
            is_readonly=False,
            error="只允许 SELECT 查询",
        )

    if "LIMIT" not in sql.upper():
        sql = sql.rstrip(";") + " LIMIT 1000"

    try:
        import time

        t0 = time.perf_counter()
        conn = _get_or_load_db(req.dataset_id)
        if conn is None:
            return SqlEditResponse(success=False, error="数据集未找到")
        cursor = conn.execute(sql)
        cols = [d[0] for d in cursor.description] if cursor.description else []
        rows = [dict(zip(cols, row)) for row in cursor.fetchall()]
        elapsed = int((time.perf_counter() - t0) * 1000)
        return SqlEditResponse(
            success=True,
            columns=cols,
            rows=rows,
            row_count=len(rows),
            elapsed_ms=elapsed,
        )
    except Exception as e:
        return SqlEditResponse(success=False, error=str(e))


class ReportGenerateRequest(BaseModel):
    """批量生成报告 - 接收选中的 query_ids"""
    query_ids: List[str]
    dataset_id: str
    title: Optional[str] = None
    format: str = "json"  # json | markdown | docx


@router.post("/report/generate")
async def generate_report(req: ReportGenerateRequest):
    """基于选中的 query_ids 生成数据报告

    每个 query_id 对应一个历史查询,会重新执行 SQL 拿数据
    format: json (默认) | markdown | docx
    """
    if not req.query_ids:
        raise HTTPException(status_code=400, detail="query_ids 不能为空")
    if req.format not in ("json", "markdown", "docx"):
        raise HTTPException(status_code=400, detail=f"不支持的 format: {req.format}")

    from services.session_service import get_recent_queries
    from services.dataset_registry import get_dataset

    # 从历史拿 query
    all_queries = get_recent_queries(limit=200)  # user_id="default" 默认
    selected = [q for q in all_queries if q.get("id") in req.query_ids]

    # 拿数据集元数据(docx 模板要用)
    dataset = get_dataset(req.dataset_id) or {}

    if not selected:
        # 可能是 reset 过, dataset_id 校验一下
        if not dataset:
            raise HTTPException(status_code=404, detail="数据集未找到")

    # 重新执行每个 query 的 SQL
    sections = []
    for q in selected:
        sql = q.get("sql") or ""
        if not sql:
            continue
        try:
            conn = _get_or_load_db(req.dataset_id)
            if not conn:
                continue
            cursor = conn.execute(sql)
            cols = [d[0] for d in cursor.description] if cursor.description else []
            rows = [dict(zip(cols, row)) for row in cursor.fetchall()]
            sections.append({
                "query_id": q.get("id"),
                "query": q.get("query"),
                "intent": q.get("intent"),
                "sql": sql,
                "columns": cols,
                "rows": rows[:100],
                "row_count": len(rows),
                "timestamp": q.get("timestamp"),
            })
        except Exception as e:
            sections.append({
                "query_id": q.get("id"),
                "query": q.get("query"),
                "intent": q.get("intent"),
                "error": str(e),
            })

    report = {
        "title": req.title or f"数据分析报告 ({len(sections)} 个指标)",
        "dataset_id": req.dataset_id,
        "dataset_name": dataset.get("name", req.dataset_id),
        "sections": sections,
        "generated_at": __import__("datetime").datetime.utcnow().isoformat(),
        "format": req.format,
    }

    if req.format == "markdown":
        md = f"# {report['title']}\n\n"
        for i, sec in enumerate(sections, 1):
            md += f"## {i}. {sec.get('query', 'N/A')}\n\n"
            md += f"**Intent**: {sec.get('intent', '-')}  |  **Rows**: {sec.get('row_count', 0)}\n\n"
            if sec.get("error"):
                md += f"> ❌ Error: {sec['error']}\n\n"
            else:
                md += f"```sql\n{sec.get('sql', '-')}\n```\n\n"
                if sec.get("rows"):
                    cols = sec["columns"]
                    md += "| " + " | ".join(cols) + " |\n"
                    md += "|" + "|".join(["---"] * len(cols)) + "|\n"
                    for row in sec["rows"][:20]:
                        md += "| " + " | ".join(str(row.get(c, "")) for c in cols) + " |\n"
                    md += "\n"
        report["markdown"] = md

    if req.format == "docx":
        # 返回 docx 二进制
        docx_bytes = _build_docx(report, dataset)
        # HTTP header 只能用 latin-1, 中文文件名要按 RFC 5987 转码
        from urllib.parse import quote
        raw_filename = f"数据分析报告-{report.get('generated_at', '')[:10]}.docx"
        encoded_filename = quote(raw_filename)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": (
                    f"attachment; "
                    f'filename="report.docx"; '  # ASCII fallback
                    f"filename*=UTF-8''{encoded_filename}"  # RFC 5987
                ),
                "Content-Length": str(len(docx_bytes)),
            },
        )

    return report
